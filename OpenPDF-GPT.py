import camelot
import json
import openai
import time
import PyPDF2
import concurrent.futures
import warnings
from retrying import retry
import pinecone
import pandas as pd
import json
from langchain.vectorstores import Pinecone
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.chains.question_answering import load_qa_chain
from langchain.chat_models import ChatOpenAI
from langchain.schema import Document
from docx import Document as DocxDocument

# API Keys
OPENAI_API_KEY = ""
PINECONE_API_KEY = ""
PINECONE_API_ENV = ''
PINECONE_INDEX = ""
openai.api_key = OPENAI_API_KEY

# Initialize Pinecone and clear the index of existing vectors
pinecone.init(api_key=PINECONE_API_KEY, environment=PINECONE_API_ENV)
index = pinecone.Index(PINECONE_INDEX)
index.delete(delete_all=True)

print(f'\n{index.describe_index_stats()}\n')

# Initialize OpenAIEmbeddings
embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)

# Initialize OpenAI LLM
llm = ChatOpenAI(model='gpt-4',temperature=0, openai_api_key=OPENAI_API_KEY)


# ------------------ AI PDF EXTRACTION FUNCTIONS ------------------

# Function to extract plain text from a PDF page
def extract_text_from_pdf_page(pdf_path, page_number):
    with open(pdf_path, 'rb') as file:
        pdf = PyPDF2.PdfReader(file)
        page = pdf.pages[page_number]
        text = page.extract_text()

    print(f"{'Plain Text':<10} {'Successfully Extracted From Page:':<30}{page_number:>5}")
    return text

# Function to extract tables from a PDF page using Camelot
def extract_tables_from_pdf_page_camelot(pdf_path, page_number):
    with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            # Read the PDF and extract tables
            tables = camelot.read_pdf(pdf_path, pages=str(page_number+1), flavor='stream')

    print(f"{'Table Data':<10} {'Successfully Extracted From Page:':<30}{page_number:>5}\n")
    return tables

# Function to get number of pages in a PDF
def get_num_pages(pdf_path):
    with open(pdf_path, 'rb') as file:
        pdf = PyPDF2.PdfReader(file)
        num_pages = len(pdf.pages)
    return num_pages

# Function to reconstruct the PDF to pass to LLM
def reconstruct_pdf(newPages):
    completion_tokens = 0
    prompt_tokens = 0
    # Function to run OpenAI API requests concurrently
    @retry(retry_on_exception=lambda e: True, wait_fixed=5000, stop_max_attempt_number=3)
    def process_page(content, page_number):
        nonlocal completion_tokens
        nonlocal prompt_tokens
        table_data = content[0]
        plain_text = content[1]

        completion = openai.ChatCompletion.create(
            model='gpt-4',
            temperature=0.2,
            messages=[
                {"role": "system", "content": "The following is extracted data from a PDF."\
                                              "First is a scan of all of the possible tables or structured text, then a simple OCR scan."\
                                              "Take both of these to build one concise page that extracts all data points and context in json"\
                                              "For long chunks of text, summarize them down and make them very brief, ignore disclaimers."},
                {"role": "user", "content": f": Table Data: {table_data}\nPlain Text: {plain_text}\nExtracted Data in JSON: "}
            ]
        )

        newJSON = completion['choices'][0]['message']['content']

        completion_tokens += completion['usage']['completion_tokens']
        prompt_tokens += completion['usage']['prompt_tokens']

        display_JSON = newJSON.replace("\n", "")
        print(f"Page {page_number+1} Extracted: {display_JSON[:120]}")

        completion = openai.ChatCompletion.create(
            model='gpt-4',
            temperature=0.2,
            messages=[
                {"role": "system", "content": "Your job is to take this raw JSON and convert it to a human readable summary."\
                                              "Break this into chunks of relevant content, separated by a newline."},
                {"role": "user", "content": f": Here is the JSON you need to summarize in human language, be verbose: {newJSON}"}
            ]
        )

        summary_for_embedding = completion['choices'][0]['message']['content']
        summary_preview = summary_for_embedding.replace("\n", "")
        print(f"Page {page_number+1} Summary: {summary_preview[:120]}")


        completion_tokens += completion['usage']['completion_tokens']
        prompt_tokens += completion['usage']['prompt_tokens']


        #print(newJSON)
        #print([page_number + 1, [newJSON, summary_for_embedding]])
        return [page_number, [newJSON, summary_for_embedding]]

    newPDF = {}

    # Create a thread pool and submit the tasks
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = [executor.submit(process_page, content, page_number) for page_number, content in enumerate(newPages)]
        for future in concurrent.futures.as_completed(futures):
            call_result = future.result()
            page_idx = call_result[0]
            page_content = call_result[1]
            newPDF[page_idx] = page_content

    
    print()
    json_document = DocxDocument()
    summary_document = DocxDocument()

    for page in sorted(newPDF):
        #display_JSON = newPDF[page][0].replace("\n", "")
        #print(f"Page {page}: {display_JSON[:120]}")

        # Add JSON data to the document
        json_document.add_heading(f'Page {page+1}', 0)
        json_document.add_heading('Extracted JSON:', level=1)
        json_document.add_paragraph(newPDF[page][0])

        # Add the summary to the document
        summary_document.add_heading(f'Page {page+1}', 0)
        summary_document.add_heading('Summary For Embedding:', level=1)
        summary_document.add_paragraph(newPDF[page][1])

        # Add a page break to the document
        json_document.add_page_break()
        summary_document.add_page_break()
        print(f"Writing Page {page}")

    json_document.save('richmond_json_test1.docx')
    summary_document.save('richmond_summary_test1.docx')

    evauluate_cost(prompt_tokens, completion_tokens)

    return newPDF                 

# Function to compute the cost of PDF processing
def evauluate_cost(prompt_tokens, completion_tokens):
    print(f"\n{'Prompt Tokens Used:':<35}{prompt_tokens:,}")  
    print(f"{'Completion Tokens Used:':<35}{completion_tokens:,}")
    print(f"\n{'Total Tokens Used:':<35}{prompt_tokens + completion_tokens:,}") 

    gpt_35_turbo_usage_cost = (prompt_tokens + completion_tokens) * 0.002 / 1000  
    print(f"\n{'GPT-3.5-Turbo Usage Cost: ':<35}${gpt_35_turbo_usage_cost:.4f}") 

    gpt_4_usage_cost = (prompt_tokens * 0.03 / 1000) + (completion_tokens * 0.06 / 1000)
    print(f"{'GPT-4 Usage Cost: ':<35}${gpt_4_usage_cost:.4f}\n") 

# Function to scrape data from PDF Page
def scrape_pages(pdf_path):
    new_pages = []
    number_of_pages = get_num_pages(pdf_path)

    for page in range(0, number_of_pages):
    #for page in range(0, 3):                                                   # <---------- CHANGE, FOR TESTING.
        page_text = extract_text_from_pdf_page(pdf_path, page)

        tables = extract_tables_from_pdf_page_camelot(pdf_path, page)
        pretty_tables = '\n'.join([json.dumps(json.loads(table.df.to_json(orient='index')), indent=4) for table in tables])

        new_pages.append([pretty_tables, page_text])
    return new_pages


# ------------------ PINECONE SEMANTIC SEARCH/INDEXING FUNCTIONS ------------------

# Function to convert files to Pinecone documents
def convert_to_pinecone_documents(newPDF):    
    vectors = []
    for i, page in enumerate(newPDF.keys()):
        chunks = newPDF[page][1].strip().split('\n')
        chunks = [x for x in chunks if x != '']
        for j, chunk in enumerate(chunks):
            print(f"Upserting Page {i+1} Chunk {j} Content: {chunk[:110]}...")
            vector = { 
                "id": f"item_{i}{j}",
                "metadata": { 
                    'page_number' : page,
                    'chunk': j},
                "values": openai.Embedding.create(input=chunk, model="text-embedding-ada-002")['data'][0]['embedding']
            }
            vectors.append(vector)
        print()

    index.upsert(vectors=vectors)
    print()

    return Pinecone.from_existing_index(PINECONE_INDEX, embeddings)

# Function to load and run the question-answering chain
def load_and_run_qa_chain(question, pdf):
    # Get vector embedding for the query
    query_vector = openai.Embedding.create(input=query,engine="text-embedding-ada-002")['data'][0]['embedding']

    # Fetch relevant documents from Pinecone
    docs = index.query(query_vector, top_k=3, include_metadata=True)

    # Calculate confidence score
    confidence_score = float(docs['matches'][0]['score']) * 100

    relevant_docs = []
    for i, doc in enumerate(docs['matches']):
        # Extract the page number and content from the relevant document
        page_number = int(doc['metadata']['page_number'])
        page_content = pdf[page_number][0]

        # Create a new doc, and make sure it is not already added for context (token purposes)
        new_doc = Document(page_content=page_content)
        if new_doc not in relevant_docs:
            relevant_docs.append(new_doc)

            display_JSON = page_content.replace("\n", "")
            print(f"Context From Page {page_number}: {display_JSON[:120]}")


    chain = load_qa_chain(llm)
    return chain.run(input_documents=relevant_docs, question=question).strip(), f"Confidence Score: {confidence_score:.4f}%"


# ------------------ MAIN FUNCTION ------------------

if __name__ == "__main__":
    start = time.time()

    pdf_path = "sample.pdf"

    newPDF = scrape_pages(pdf_path)
    finalPDF = reconstruct_pdf(newPDF)
    pinecone_docs = convert_to_pinecone_documents(finalPDF)
    
    #pinecone_docs = Pinecone.from_existing_index(PINECONE_INDEX, embeddings)

    queries= ["Sample Question?",
              "Sample Question?",
              "Sample Question?",
              "Sample Question?",
              "Sample Question?"",
              "Sample Question?",
              "Sample Question?",
              "Sample Question?"]

    for query in queries:
        answer, confidence = load_and_run_qa_chain(question=query,pdf=finalPDF)
        print(f'\n{confidence}\n{answer}\n')
    
    end = time.time()
    elapsed_time = end - start
    print(f"Elapsed time: {elapsed_time:.2f} seconds")
